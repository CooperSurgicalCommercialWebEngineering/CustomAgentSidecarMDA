from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent
DOCX_PATH = OUT / "HR-Management-App-Guide-Entra-App-Registration.docx"

BLUE = "17365D"
TEAL = "0F6CBD"
LIGHT_BLUE = "EAF2F8"
LIGHT_GOLD = "FFF4CE"
LIGHT_GRAY = "F3F4F6"
WHITE = "FFFFFF"
DARK = RGBColor(31, 41, 55)
GRAY = RGBColor(91, 101, 115)

ENVIRONMENT_URL = "https://carremacodeapps.crm.dynamics.com"
REDIRECT_URI = f"{ENVIRONMENT_URL}/WebResources/maftagsc_/copilot/authRedirect.html"
ENVIRONMENT_ID = "f9b87f8b-0abf-e629-affb-b13195d1ed14"
AGENT_SCHEMA = "cr0b1_HRMgmtClassic"
DELEGATED_SCOPE = "https://api.powerplatform.com/CopilotStudio.Copilots.Invoke"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, color=WHITE)
        shade(cell, BLUE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], str(value))
            if row_index % 2:
                shade(cells[index], LIGHT_GRAY)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for r2 in p2.runs:
        r2.font.name = "Arial"
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_step(doc, number, title, actions):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run(f"{number}. {title}")
    for action in actions:
        para = doc.add_paragraph(style="List Number")
        para.add_run(action)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Title", 25, BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 12.5, TEAL),
        ("Heading 3", 11, BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("HR Agent Sidecar  |  Entra app registration guide  |  Internal")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = GRAY


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([run_properties, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_document():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "HR Management App Guide — Microsoft Entra App Registration"
    doc.core_properties.subject = "Delegated authentication setup for a Copilot Studio agent in a model-driven app"
    doc.core_properties.author = "HR Agent Sidecar project team"
    doc.core_properties.comments = "Generated project documentation"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(62)
    r = p.add_run("HR AGENT SIDECAR")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(TEAL)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Microsoft Entra App Registration Guide")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Delegated Copilot Studio access from the HR Management model-driven app")
    sr.font.name = "Arial"
    sr.font.size = Pt(14)
    sr.font.color.rgb = GRAY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Version 1.0  |  July 9, 2026  |  Administrator and developer procedure")
    mr.font.name = "Arial"
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_callout(
        doc,
        "Objective",
        "Create a single-tenant browser client registration that lets the HR Management side pane call the published Copilot Studio agent on behalf of the signed-in Microsoft user. This preserves user-scoped SharePoint access and removes the need to treat a Direct Line token as user authentication.",
    )
    add_callout(
        doc,
        "Security outcome",
        "The browser receives only delegated, short-lived access tokens. No client secret is created or embedded. Copilot Studio remains configured to Authenticate with Microsoft, and SharePoint continues to enforce the signed-in user's existing permissions.",
        LIGHT_GOLD,
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("1. Configuration worksheet", level=1)
    add_table(
        doc,
        ["Setting", "Required value"],
        [
            ["Registration name", "HR Management App Guide Sidecar - Dev"],
            ["Supported account type", "Accounts in this organizational directory only (single tenant)"],
            ["Platform", "Single-page application (SPA)"],
            ["SPA redirect URI", REDIRECT_URI],
            ["Delegated API permission", "Power Platform API > CopilotStudio > CopilotStudio.Copilots.Invoke"],
            ["Runtime scope", DELEGATED_SCOPE],
            ["Environment ID", ENVIRONMENT_ID],
            ["Agent schema name", AGENT_SCHEMA],
            ["Application (client) ID", "Record after registration: __________________________________"],
            ["Directory (tenant) ID", "Record after registration: __________________________________"],
        ],
        [2.25, 4.65],
    )

    doc.add_heading("2. Prerequisites", level=1)
    for text in [
        "A Microsoft Entra account permitted to create app registrations.",
        "A tenant administrator available to grant tenant-wide consent if your role cannot do so.",
        "Access to the published HR Management App Guide agent in Copilot Studio.",
        "The target Dataverse organization URL and the agent metadata values shown above.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("3. Registration procedure", level=1)
    add_step(doc, 1, "Open Microsoft Entra admin center", [
        "Go to https://entra.microsoft.com and sign in.",
        "Confirm the correct tenant is selected.",
        "Navigate to Identity > Applications > App registrations.",
        "Select New registration.",
    ])
    add_step(doc, 2, "Register the browser client", [
        "Enter the name HR Management App Guide Sidecar - Dev.",
        "Select Accounts in this organizational directory only (single tenant).",
        "Under Redirect URI, select Single-page application (SPA).",
        f"Enter {REDIRECT_URI} exactly, then select Register.",
    ])
    add_callout(doc, "Why SPA?", "The side pane is browser-hosted. SPA registration enables authorization-code flow with PKCE. Do not register this client as a confidential Web application.")

    add_step(doc, 3, "Record non-secret identifiers", [
        "On Overview, copy the Application (client) ID into the worksheet.",
        "Copy the Directory (tenant) ID into the worksheet.",
        "Treat both as configuration identifiers, not passwords. Never substitute the Copilot Studio Agent App ID for this client ID.",
    ])
    add_step(doc, 4, "Verify authentication settings", [
        "Open Manage > Authentication.",
        f"Confirm a Single-page application platform exists with {REDIRECT_URI}.",
        "Leave Access tokens and ID tokens under Implicit grant and hybrid flows cleared.",
        "Under Advanced settings, set Allow public client flows to Yes, as required by the Microsoft reference architecture.",
        "Select Save.",
    ])
    add_step(doc, 5, "Add delegated Copilot Studio permission", [
        "Open Manage > API permissions and select Add a permission.",
        "Select APIs my organization uses and search for Power Platform API.",
        "Select Delegated permissions.",
        "Expand CopilotStudio and select CopilotStudio.Copilots.Invoke.",
        "Select Add permissions.",
    ])
    add_step(doc, 6, "Grant administrator consent", [
        "On API permissions, select Grant admin consent for the tenant.",
        "Confirm the prompt.",
        "Verify the permission is Delegated and its status is Granted for the tenant.",
    ])
    add_callout(doc, "Do not create a client secret", "A browser cannot safely keep a client secret. The client uses delegated authentication, PKCE, and short-lived access tokens. Certificates & secrets should remain empty for this registration.", LIGHT_GOLD)

    add_step(doc, 7, "Verify the Copilot Studio agent", [
        "Open HR Management App Guide in Copilot Studio.",
        "Under Settings > Security > Authentication, keep Authenticate with Microsoft selected.",
        "Under Settings > Advanced > Metadata, confirm the Environment ID and Schema name match the worksheet.",
        "From the agent overview, share chat access with the intended HR users or an Entra security group containing the test user.",
        "Publish the agent if its configuration changed.",
    ])

    doc.add_heading("4. What the four values are used for", level=1)
    add_table(
        doc,
        ["Value", "Runtime use", "Secret?"],
        [
            ["Application (client) ID", "Identifies the browser client when MSAL requests a delegated Power Platform API token.", "No"],
            ["Directory (tenant) ID", "Builds the tenant-specific Microsoft identity authority and keeps sign-in single-tenant.", "No"],
            ["Environment ID", "Tells CopilotStudioClient which Power Platform environment hosts the agent.", "No"],
            ["Agent schema name", "Identifies the published agent supplied as agentIdentifier to CopilotStudioClient.", "No"],
        ],
        [2.0, 4.15, 0.75],
    )

    doc.add_heading("5. Developer handoff", level=1)
    p = doc.add_paragraph("Provide the developer these four non-secret values:")
    for text in [
        "Application (client) ID from the new registration",
        "Directory (tenant) ID from the new registration",
        f"Environment ID: {ENVIRONMENT_ID}",
        f"Agent schema name: {AGENT_SCHEMA}",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph("The implementation uses them as follows:")
    code_lines = [
        "MSAL clientId             = <Application (client) ID>",
        "MSAL authority            = https://login.microsoftonline.com/<Directory (tenant) ID>",
        f"MSAL redirectUri          = {REDIRECT_URI}",
        f"Requested delegated scope = {DELEGATED_SCOPE}",
        f"Agents SDK environmentId  = {ENVIRONMENT_ID}",
        f"Agents SDK agentIdentifier= {AGENT_SCHEMA}",
    ]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "F8FAFC")
    cell.text = ""
    for index, line in enumerate(code_lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)

    doc.add_heading("6. Validation checklist", level=1)
    for text in [
        "The app registration is single tenant.",
        "The redirect URI is registered under SPA and exactly matches the dedicated Dataverse authentication redirect web resource.",
        "CopilotStudio.Copilots.Invoke is a delegated permission and tenant consent is granted.",
        "No client secret exists for the browser client.",
        "The agent remains set to Authenticate with Microsoft and is published.",
        "The test user has permission to chat with the agent.",
        "The test user has access to the intended SharePoint knowledge files.",
        "First use may show a consent or sign-in popup; subsequent token acquisition should normally be silent.",
        "A user without SharePoint access does not receive protected SharePoint content.",
    ]:
        doc.add_paragraph(f"☐ {text}")

    doc.add_heading("7. Troubleshooting", level=1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Action"],
        [
            ["AADSTS redirect URI mismatch", "The exact runtime redirect URI is not registered as SPA.", "Compare scheme, hostname, path, and trailing slash; use the worksheet value exactly."],
            ["Consent required or permission missing", "Delegated permission lacks administrator consent.", "Grant tenant consent and verify the permission status."],
            ["No account available for silent token acquisition", "The browser client has no cached account.", "Use interactive popup sign-in once, then retry silent acquisition."],
            ["Agent invocation returns forbidden", "Agent is not shared with the user or the wrong tenant/client is used.", "Verify sharing, tenant ID, client ID, and published agent metadata."],
            ["SharePoint answer is unavailable", "The user lacks SharePoint access, the source is restricted, or agent configuration is unpublished.", "Verify the user's direct SharePoint access and republish the agent if needed."],
        ],
        [2.05, 2.25, 2.6],
    )

    doc.add_heading("8. Official references", level=1)
    references = [
        ("Microsoft model-driven app and Copilot Studio reference architecture", "https://learn.microsoft.com/dynamics365/guidance/reference-architectures/custom-copilot-agent-dynamics-365-power-apps"),
        ("Configure user authentication in Copilot Studio", "https://learn.microsoft.com/microsoft-copilot-studio/configuration-end-user-authentication"),
        ("Add SharePoint as a knowledge source", "https://learn.microsoft.com/microsoft-copilot-studio/knowledge-add-sharepoint"),
        ("Integrate with web or native apps using Microsoft 365 Agents SDK", "https://learn.microsoft.com/microsoft-copilot-studio/publication-integrate-web-or-native-app-m365-agents-sdk"),
    ]
    for label, url in references:
        paragraph = doc.add_paragraph(style="List Bullet")
        add_hyperlink(paragraph, label, url)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
